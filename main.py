"""
搜索评测 Web 应用 - FastAPI 后端
"""

import os
import csv
import json
import time
import sqlite3
import asyncio
import uuid
from io import BytesIO, StringIO
from datetime import datetime
from typing import List

import aiosqlite
import openpyxl
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
from pydantic import BaseModel

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(APP_DIR, "data.db")
SCREENSHOTS_DIR = os.path.join(APP_DIR, "static", "screenshots")
STATIC_DIR = os.path.join(APP_DIR, "static")

os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

app = FastAPI(title="搜索评测系统")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# Import progress tracking
import_tasks = {}  # task_id -> {status, progress, total, imported, error, result}

# ============================================================
# Database
# ============================================================

COLUMNS = [
    "id", "query", "query_length", "business_category", "search_count",
    "is_short_query", "frequency_level", "original_group", "search_result",
    "text_screenshot_url", "ai_screenshot_url",
    "relevance", "accuracy", "ranking", "diversity",
    "quality_threshold", "copy_recommendation", "click_desire",
    "vs_text_search", "ai_search_score",
    "created_at", "updated_at"
]

EXCEL_COLUMNS = [
    "query", "query_length", "business_category", "search_count",
    "is_short_query", "frequency_level", "original_group", "search_result",
    "expected_screenshot", "expected_result",
    "text_screenshot_url", "ai_screenshot_url",
    "relevance", "accuracy", "ranking", "diversity",
    "quality_threshold", "copy_recommendation", "click_desire",
    "vs_text_search", "ai_search_score",
]

EXCEL_HEADERS = [
    "query", "query_length", "业务分类", "搜索次数",
    "是否短query", "频率层级", "原始分组", "搜索结果",
    "文本搜截图", "AI搜截图",
    "相关性", "准确性", "排序合理性", "多样性",
    "质量门槛（与文本搜索对比）", "文案/推荐语", "点击欲望",
    "与文本搜相比", "文本搜索综合分", "ai搜索综合分",
    "业务分类"
]

# Map Chinese Excel header -> DB field name
HEADER_TO_FIELD = {
    "query": "query",
    "query_length": "query_length",
    "业务分类": "business_category",
    "搜索次数": "search_count",
    "是否短query": "is_short_query",
    "频率层级": "frequency_level",
    "原始分组": "original_group",
    "期望结果截图": "expected_screenshot",
    "期望搜索结果": "expected_result",
    "搜索结果": "search_result",
    "文本搜截图": "text_screenshot_url",
    "AI搜截图": "ai_screenshot_url",
    "ai搜截图": "ai_screenshot_url",
    "相关性": "relevance",
    "准确性": "accuracy",
    "排序合理性": "ranking",
    "多样性": "diversity",
    "质量门槛": "quality_threshold",
    "质量门槛（与文本搜索对比）": "quality_threshold",
    "文案/推荐语": "copy_recommendation",
    "点击欲望": "click_desire",
    "与文本搜相比": "vs_text_search",
    "VS文本搜": "vs_text_search",
    "vs文本搜": "vs_text_search",
    "VS文本搜索": "vs_text_search",
    "与文本搜对比": "vs_text_search",
    "文本搜索综合分": None,
    "ai搜索综合分": "ai_search_score",
    "AI搜索综合分": "ai_search_score",
    "业务分类2": "business_category",
}


# 关键词表：只要表头（归一化后）包含某关键词即自动匹配到对应字段。
# 用最长关键词优先，避免歧义（如 query_length 优先于 query）。
FIELD_KEYWORDS = {
    "query_length": ["querylength", "长度", "字数"],
    "business_category": ["业务分类", "业务类别", "业务类型", "业务"],
    "search_count": ["搜索次数", "搜索量", "查询次数", "次数"],
    "is_short_query": ["是否短query", "短query", "query层级", "长短query", "长短"],
    "frequency_level": ["频率层级", "频率", "频次"],
    "original_group": ["原始分组", "分组"],
    "expected_screenshot": ["期望搜索结果截图", "期望结果截图", "期望截图", "理想截图"],
    "expected_result": ["期望搜索结果", "期望结果", "理想地图", "理想结果", "期望"],
    "search_result": ["搜索结果说明", "实际搜索结果", "搜索结果", "实际结果"],
    "text_screenshot_url": ["文本搜截图", "文本搜索截图", "文本截图"],
    "ai_screenshot_url": ["ai搜截图", "ai搜索截图", "ai截图"],
    "relevance": ["相关性", "相关"],
    "accuracy": ["准确性", "准确"],
    "ranking": ["排序合理性", "排序"],
    "diversity": ["多样性", "多样"],
    "quality_threshold": ["质量门槛", "质量"],
    "copy_recommendation": ["文案推荐语", "推荐语", "文案"],
    "click_desire": ["点击欲望", "点击意愿", "点击"],
    "vs_text_search": ["vs文本", "与文本搜相比", "与文本搜对比", "文本搜对比", "文本搜相比", "对比文本搜"],
    "ai_search_score": ["ai搜索综合分", "ai综合分", "ai搜综合分", "ai评分"],
    "query": ["query", "查询词", "搜索词"],
}
_KEYWORD_LIST = sorted(
    [(kw, field) for field, kws in FIELD_KEYWORDS.items() for kw in kws],
    key=lambda x: -len(x[0]),
)

# 表头映射到 None 的列 = 已知但主动忽略（不算未识别）
_IGNORE = "__IGNORE__"


def _norm_header(s):
    """归一化表头：小写、去空格与常见标点/括号，便于容错匹配。"""
    s = str(s or "").lower()
    for ch in [" ", "\t", "\n", "\r", "　", "（", "）", "(", ")", "【", "】",
               "[", "]", "、", "，", ",", "/", "·", "-", "_", "：", ":", ".", "。"]:
        s = s.replace(ch, "")
    return s


def match_header(header, dynamic_mapping):
    """智能匹配表头到字段。返回字段名 / _IGNORE / None（未识别）。"""
    import difflib
    if not header:
        return None
    nh = _norm_header(header)
    if not nh:
        return None
    # 归一化后的映射表（内置 + 自定义列）
    norm_map = {}
    for k, v in dynamic_mapping.items():
        nk = _norm_header(k)
        if nk and nk not in norm_map:
            norm_map[nk] = v
    # 1) 归一化精确匹配
    if nh in norm_map:
        return norm_map[nh] if norm_map[nh] is not None else _IGNORE
    # 2) 关键词匹配（内置字段，最长关键词优先）
    for kw, field in _KEYWORD_LIST:
        if kw in nh:
            return field
    # 3) 与映射键互为子串（主要覆盖自定义列标签）
    for nk, v in norm_map.items():
        if v and (nk in nh or nh in nk):
            return v
    # 4) 相似度兜底
    best, best_ratio = None, 0.0
    for nk, v in norm_map.items():
        if not v:
            continue
        r = difflib.SequenceMatcher(None, nh, nk).ratio()
        if r > best_ratio:
            best, best_ratio = v, r
    if best and best_ratio >= 0.72:
        return best
    return None


# ============================================================
# Custom Columns
# ============================================================

CUSTOM_COLUMNS_PATH = os.path.join(APP_DIR, "custom_columns.json")


def load_custom_columns(test_set_id=None):
    """Load custom column definitions: [{field, label, type}]"""
    if test_set_id:
        path = os.path.join(APP_DIR, f"custom_columns_{test_set_id}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    # Fallback to global (for backward compat)
    if os.path.exists(CUSTOM_COLUMNS_PATH):
        with open(CUSTOM_COLUMNS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_custom_columns(cols, test_set_id=None):
    if test_set_id:
        path = os.path.join(APP_DIR, f"custom_columns_{test_set_id}.json")
    else:
        path = CUSTOM_COLUMNS_PATH
    with open(path, "w", encoding="utf-8") as f:
        json.dump(cols, f, ensure_ascii=False, indent=2)


REPORT_NOTES_FIELDS = ["conclusion", "priorities", "stats", "badcase"]


DEFAULT_BADCASE = """4.1 问题汇总（算法填写可能原因与解决办法）
体感 / 问题：
- 首图质量太低 → 排序
- 总体质量偏低 → 召回、排序
- 缺图漏图，推不满9张图
- 内容理解不准
- query未改写 → 意图理解
- 推图不准确 → 意图理解、召回

问题一：质量偏低
1. 后缀语义处理：乐园地图→地图、和朋友一起→多人、新地图→近期热门图
2. 限定搜索范围：下一期上线分类器，缩小搜索范围至2w带地图信息的热门推荐图（游玩量+点赞比双阈值保障），先在AI搜索场景测试效果
3. 排序策略优化：6张结果中最多出1-2张少于1万游玩量的图（新图扶持机制，仅1/6~2/6流量测试）
问题二：内容理解不准确
1. 增加地图解析字段：优先解析"最大出生点"字段实现人数识别，新增颜色字段；根据AI搜索日志推导缺失标签，制定二次全量打标字段清单
2. 完善AI知识库：总结热门榜单高频玩法描述词，支持后端增量更新
问题三：返回结果缺图漏图
1. 提单增加返回地图请求数量（5张→10张），客户端做兜底填充
问题四：query纠错和改写
1. 持续收集改写案例，后续安排LLM模型改写优化（优先级后排）

4.2 场景限定
1. 适合两个人一起玩的乐园地图（频率43）｜期望：排行榜双人《卷纸双蛋（双人）》《双人骑单车》｜问题：首图质量太低、总体质量偏低（尤其"我的偏好"）｜原因：乐园地图后处理
2. 适合两个人一起玩的地图（41）｜总体质量良好（对照组）
3. 适合和朋友一起玩的乐园地图（21）｜期望：多人热门地图｜问题：总体质量偏低（尤其"我的偏好"）｜原因：AI识别
4. 有趣好玩的乐园地图（26）｜期望：热门排行前几｜问题：首图质量低、总体质量偏低、推不满9张｜原因：限定搜索范围、融合搜场景用小范围地图
5. 好玩的新地图推荐（5）｜问题：是否可全部推近期热门地图
6. 色彩丰富风景美丽适合休闲玩的跑酷图（1）｜问题：召回"近期热门"相关性较低

4.2 场景限定-观赏（1、质量不高：文本理解需加强 2、缺图漏图）
1. 有氛围感、花海颜色很暖的地图（4）｜期望：《十里桃花（运镜）》《夕岚花镜》｜问题：相关性不高｜原因：文本理解、召回、打标（人数/颜色）
2. 全自动、手都不用动的音乐图（51）｜期望：音乐《打火机》《一点点》、RUDE/Whiplash全自动观赏｜问题：推不满9张、相关性不高｜原因：打标
3. 想要和"四片叶子"一样好看的观赏图（1）｜期望：《云野西瓜》《云野荔枝》《铃兰花序》｜问题：总体质量偏低、缺热门观赏图｜原因：文本理解

4.3 场景限定-难度（总体质量不高）
1. 很难的地图（19）｜问题：首图质量低、"近期热门"质量低、总体质量偏低｜原因：游玩量/点赞比控制，9张里可有1-2张偏低（7天增量）
2. 有点难的地图（6）｜问题：总体质量偏低
3. 轻松有趣的竞速地图（1）｜问题：相关性不足、总体质量偏低

4.4 场景限定-恐怖（缺图漏图很严重）
1. 来一个恐怖地图（4）｜问题：缺图漏图严重｜原因：服务端请求数量增加（提单）
2. 让人害怕的恐怖主题地图（107）｜问题：首图质量低、缺图漏图、总体质量偏低
3. 超级恐怖且收藏多玩的人也多（1）｜问题：不满足query条件（收藏人多）、缺图漏图
4. 非常恐怖的地图（1）｜问题：相关性不足、总体质量偏低、缺图漏图

4.5 玩法描述（1、推图不准确 2、地图质量不高 3、改写失败）
1. 我和我的小仓鼠（40）｜问题：召回query匹配不准确
2. 过生日找生日快乐可自己开店的地图（1）｜问题：生日推了恐怖图、总体质量不高｜原因：开店=步行街（官方tag），先依赖内容理解
3. 美观一关妈妈年龄加一（1）｜问题：改写失败，没推出理想图｜原因：知识库（热门地图高频词"每过一关"）
4. 想玩拼斗真正的可以拼（1）｜期望：《拼豆实景工坊》｜问题：相关性不高、改写失败
5. 宝宝乐园摔下来相关（1）｜问题：没推出理想图
6. 寻找以前玩过的擂台（7）｜期望：推热门图｜问题：缺图漏图
7. 体验载具的地图（7）｜问题：缺图漏图｜原因：黑产图被UGC限流
8. 1v2咸鱼/擂台（1）｜问题：缺图漏图｜原因：黑产图，作者被拉黑
9. 很难的特性图（4）｜问题：推不出正版

4.6 疑问句&商业化（特殊，考虑拓展该部分回复内容）
1. 可用乐园B换的枪战地图（1）
2. 很多要充钱才能获得载具的地图（1）｜原因：联网可小规模尝试
3. 中奖概率倍儿高（抽奖）（1）
4. 可领笑脸币买皮肤的游戏（1）
5. 可抽新赛季皮肤的地图（1）
疑问query可制定专属回复："地图马上出现！~" / "还在跟脑细胞斗争中。。"
- 怎么才能找到店铺？（16）
- 可以开了吗？（26）
（注：各case截图见原POPO文档）"""


def load_report_notes(test_set_id=None):
    """Load per-test-set free-text report notes."""
    tid = test_set_id or 0
    path = os.path.join(APP_DIR, f"report_notes_{tid}.json")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {}
    notes = {k: data.get(k, "") for k in REPORT_NOTES_FIELDS}
    return notes


def save_report_notes(notes, test_set_id=None):
    tid = test_set_id or 0
    path = os.path.join(APP_DIR, f"report_notes_{tid}.json")
    clean = {k: str(notes.get(k, "") or "") for k in REPORT_NOTES_FIELDS}
    with open(path, "w", encoding="utf-8") as f:
        json.dump(clean, f, ensure_ascii=False, indent=2)
    return clean


KNOWLEDGE_PATH = os.path.join(APP_DIR, "knowledge.json")
KNOWLEDGE_FIELDS = ["intro", "ai_def", "flow", "scoring", "badcase_std"]


def load_knowledge():
    if os.path.exists(KNOWLEDGE_PATH):
        with open(KNOWLEDGE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {}
    return {k: data.get(k, "") for k in KNOWLEDGE_FIELDS}


def save_knowledge(notes):
    clean = {k: str(notes.get(k, "") or "") for k in KNOWLEDGE_FIELDS}
    with open(KNOWLEDGE_PATH, "w", encoding="utf-8") as f:
        json.dump(clean, f, ensure_ascii=False, indent=2)
    return clean


def get_dynamic_header_mapping(test_set_id=None):
    """Return HEADER_TO_FIELD merged with custom columns"""
    mapping = dict(HEADER_TO_FIELD)
    for col in load_custom_columns(test_set_id):
        mapping[col["label"]] = col["field"]
    return mapping


def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS queries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            query TEXT,
            query_length INTEGER,
            business_category TEXT,
            search_count TEXT,
            is_short_query TEXT,
            frequency_level TEXT,
            original_group TEXT,
            search_result TEXT,
            expected_screenshot TEXT,
            expected_result TEXT,
            text_screenshot_url TEXT,
            ai_screenshot_url TEXT,
            relevance REAL,
            accuracy REAL,
            ranking REAL,
            diversity REAL,
            quality_threshold REAL,
            copy_recommendation TEXT,
            click_desire REAL,
            vs_text_search TEXT,
            ai_search_score REAL,
            created_at TEXT,
            updated_at TEXT
        )
    """)
    # Test sets table
    conn.execute("""
        CREATE TABLE IF NOT EXISTS test_sets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            created_at TEXT
        )
    """)
    # Ensure test_set_id column exists
    existing = {row[1] for row in conn.execute("PRAGMA table_info(queries)").fetchall()}
    if "test_set_id" not in existing:
        conn.execute("ALTER TABLE queries ADD COLUMN test_set_id INTEGER DEFAULT 1")
    if "expected_screenshot" not in existing:
        conn.execute("ALTER TABLE queries ADD COLUMN expected_screenshot TEXT")
    if "expected_result" not in existing:
        conn.execute("ALTER TABLE queries ADD COLUMN expected_result TEXT")
    # test_sets 标注状态：labeling(标注中) / locked(已锁定只读)
    ts_cols = {row[1] for row in conn.execute("PRAGMA table_info(test_sets)").fetchall()}
    if "status" not in ts_cols:
        conn.execute("ALTER TABLE test_sets ADD COLUMN status TEXT DEFAULT 'labeling'")
    # Ensure at least one test set exists
    count = conn.execute("SELECT COUNT(*) FROM test_sets").fetchone()[0]
    if count == 0:
        conn.execute("INSERT INTO test_sets (id, name, created_at) VALUES (1, '测试集 1', datetime('now'))")
    # Add custom columns to DB if missing
    existing = {row[1] for row in conn.execute("PRAGMA table_info(queries)").fetchall()}
    for col in load_custom_columns():
        if col["field"] not in existing:
            col_type = "REAL" if col.get("type") == "number" else "TEXT"
            conn.execute(f'ALTER TABLE queries ADD COLUMN {col["field"]} {col_type}')
    conn.commit()
    conn.close()


init_db()


async def get_db():
    db = await aiosqlite.connect(DB_PATH)
    db.row_factory = aiosqlite.Row
    return db


# ============================================================
# WebSocket Manager
# ============================================================

class ConnectionManager:
    def __init__(self):
        self.connections: List[WebSocket] = []

    async def connect(self, ws: WebSocket):
        await ws.accept()
        self.connections.append(ws)

    def disconnect(self, ws: WebSocket):
        self.connections.remove(ws)

    async def broadcast(self, message: dict):
        for ws in self.connections[:]:
            try:
                await ws.send_json(message)
            except Exception:
                self.connections.remove(ws)


manager = ConnectionManager()


@app.websocket("/ws")
async def websocket_endpoint(ws: WebSocket):
    await manager.connect(ws)
    try:
        while True:
            await ws.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(ws)


# ============================================================
# API Routes
# ============================================================

from fastapi import Query as QueryParam


@app.get("/api/test-sets")
async def get_test_sets():
    db = await get_db()
    cursor = await db.execute("SELECT * FROM test_sets ORDER BY id")
    rows = await cursor.fetchall()
    await db.close()
    return [dict(row) for row in rows]


@app.post("/api/test-sets")
async def create_test_set(body: dict):
    name = body.get("name", "").strip()
    copy_from = body.get("copy_from")  # test_set_id to copy queries from
    if not name:
        raise HTTPException(400, "名称不能为空")

    db = await get_db()
    now = datetime.now().isoformat()
    cursor = await db.execute("INSERT INTO test_sets (name, created_at) VALUES (?, ?)", [name, now])
    new_id = cursor.lastrowid

    if copy_from:
        # Copy queries but clear scoring/screenshot fields
        src_cursor = await db.execute("SELECT * FROM queries WHERE test_set_id = ?", [copy_from])
        src_rows = [dict(r) for r in await src_cursor.fetchall()]

        score_fields = ['relevance', 'accuracy', 'ranking', 'diversity', 'quality_threshold',
                        'copy_recommendation', 'click_desire', 'vs_text_search', 'ai_search_score',
                        'text_screenshot_url', 'ai_screenshot_url', 'search_result',
                        'expected_result', 'expected_screenshot']
        custom_cols = load_custom_columns(copy_from)
        score_fields += [c["field"] for c in custom_cols]

        for row in src_rows:
            row.pop("id", None)
            row["test_set_id"] = new_id
            row["created_at"] = now
            row["updated_at"] = now
            for sf in score_fields:
                row[sf] = None

            fields = [k for k in row.keys()]
            placeholders = ", ".join(["?"] * len(fields))
            field_names = ", ".join(fields)
            await db.execute(f"INSERT INTO queries ({field_names}) VALUES ({placeholders})", list(row.values()))

    await db.commit()
    await db.close()
    return {"id": new_id, "name": name, "created_at": now}


@app.put("/api/test-sets/{set_id}")
async def rename_test_set(set_id: int, body: dict):
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(400, "名称不能为空")
    db = await get_db()
    await db.execute("UPDATE test_sets SET name = ? WHERE id = ?", [name, set_id])
    await db.commit()
    await db.close()
    return {"ok": True}


@app.delete("/api/test-sets/{set_id}")
async def delete_test_set(set_id: int):
    db = await get_db()
    count = await db.execute("SELECT COUNT(*) FROM test_sets")
    total = (await count.fetchone())[0]
    if total <= 1:
        await db.close()
        raise HTTPException(400, "至少保留一个测试集")
    await db.execute("DELETE FROM queries WHERE test_set_id = ?", [set_id])
    await db.execute("DELETE FROM test_sets WHERE id = ?", [set_id])
    await db.commit()
    await db.close()
    return {"ok": True}


@app.get("/api/queries")
async def get_queries(test_set_id: int = QueryParam(default=None)):
    db = await get_db()
    if test_set_id:
        cursor = await db.execute("SELECT * FROM queries WHERE test_set_id = ? ORDER BY id", [test_set_id])
    else:
        cursor = await db.execute("SELECT * FROM queries ORDER BY id")
    rows = await cursor.fetchall()
    await db.close()
    return [dict(row) for row in rows]


class QueryUpdate(BaseModel):
    model_config = {"extra": "allow"}


@app.put("/api/test-sets/{set_id}/status")
async def set_test_set_status(set_id: int, body: dict):
    status = body.get("status", "").strip()
    if status not in ("labeling", "locked"):
        raise HTTPException(400, "status 必须是 labeling 或 locked")
    db = await get_db()
    await db.execute("UPDATE test_sets SET status = ? WHERE id = ?", [status, set_id])
    await db.commit()
    await db.close()
    return {"ok": True, "status": status}


async def _ensure_unlocked(db, query_id):
    """测试集处于 locked 时抛 423，禁止任何写入（含直接上传的结果）。"""
    row0 = await (await db.execute("SELECT test_set_id FROM queries WHERE id = ?", [query_id])).fetchone()
    if row0 is not None:
        ts = await (await db.execute("SELECT status FROM test_sets WHERE id = ?", [row0["test_set_id"]])).fetchone()
        if ts is not None and ts["status"] == "locked":
            await db.close()
            raise HTTPException(423, "测试集已锁定，请先「解锁编辑」")


@app.put("/api/queries/{query_id}")
async def update_query(query_id: int, data: QueryUpdate):
    updates = data.model_extra or {}
    if not updates:
        raise HTTPException(400, "No fields to update")

    db = await get_db()
    # 锁定校验：测试集 locked 时禁止编辑
    await _ensure_unlocked(db, query_id)

    updates["updated_at"] = datetime.now().isoformat()
    set_clause = ", ".join(f"{k} = ?" for k in updates.keys())
    values = list(updates.values()) + [query_id]

    await db.execute(f"UPDATE queries SET {set_clause} WHERE id = ?", values)
    await db.commit()

    cursor = await db.execute("SELECT * FROM queries WHERE id = ?", [query_id])
    row = await cursor.fetchone()
    await db.close()

    if row is None:
        raise HTTPException(404, "Query not found")
    return dict(row)


@app.post("/api/upload-screenshot")
async def upload_screenshot(query_id: int, type: str, file: UploadFile = File(...)):
    if type not in ("text", "ai"):
        raise HTTPException(400, "type must be 'text' or 'ai'")

    db = await get_db()
    await _ensure_unlocked(db, query_id)

    filename = f"q{query_id}_{type}_{int(datetime.now().timestamp())}.jpg"
    filepath = os.path.join(SCREENSHOTS_DIR, filename)

    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    url = f"/static/screenshots/{filename}"
    col = "text_screenshot_url" if type == "text" else "ai_screenshot_url"

    await db.execute(
        f"UPDATE queries SET {col} = ?, updated_at = ? WHERE id = ?",
        [url, datetime.now().isoformat(), query_id]
    )
    await db.commit()
    await db.close()

    await manager.broadcast({
        "type": "screenshot_update",
        "query_id": query_id,
        "screenshot_type": type,
        "url": url
    })

    return {"url": url, "query_id": query_id}


@app.post("/api/upload-field-image")
async def upload_field_image(query_id: int, field: str, file: UploadFile = File(...)):
    import re
    ALLOWED_IMG_FIELDS = {"expected_screenshot", "text_screenshot_url", "ai_screenshot_url"}
    if not (re.match(r'^custom_\w+$', field) or field in ALLOWED_IMG_FIELDS):
        raise HTTPException(400, "非法字段")

    db = await get_db()
    await _ensure_unlocked(db, query_id)
    await db.close()

    safe = re.sub(r'[^0-9A-Za-z_]', '', field) or 'img'
    filename = f"q{query_id}_{safe}_{int(datetime.now().timestamp())}.jpg"
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    url = f"/static/screenshots/{filename}"
    return {"url": url, "query_id": query_id, "field": field}


@app.post("/api/import")
async def import_excel(file: UploadFile = File(...), skip_unknown: int = 0, test_set_id: int = 1):
    import re
    import zipfile
    import tempfile
    from xml.etree import ElementTree as ET

    # Stream upload to temp file (don't hold 177MB in RAM)
    task_id = str(uuid.uuid4())[:8]
    import_tasks[task_id] = {"status": "uploading", "progress": 0, "total": 0, "imported": 0, "error": None, "result": None}

    filename_l = (file.filename or "").lower()
    is_csv = filename_l.endswith(".csv")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.csv' if is_csv else '.xlsx')
    try:
        while chunk := await file.read(1024 * 1024):  # 1MB chunks
            tmp.write(chunk)
        tmp.close()
        tmp_path = tmp.name
    except Exception as e:
        import_tasks[task_id]["status"] = "error"
        import_tasks[task_id]["error"] = str(e)
        return {"task_id": task_id, "error": str(e)}

    import_tasks[task_id]["status"] = "parsing"

    name_to_zip_path = {}
    header_row = []
    total_rows = 0

    if is_csv:
        # --- CSV：读取表头并统计行数（无内嵌图片） ---
        csv_rows = _read_csv_rows(tmp_path)
        if csv_rows:
            header_row = [str(c).strip() if c else '' for c in csv_rows[0]]
            total_rows = max(0, len(csv_rows) - 1)
    else:
        # --- Extract DISPIMG name->zip_path mapping (no image bytes in memory) ---
        try:
            zf = zipfile.ZipFile(tmp_path)
            name_to_rid = {}
            if 'xl/cellimages.xml' in zf.namelist():
                ci_xml = zf.read('xl/cellimages.xml')
                root = ET.fromstring(ci_xml)
                ns = {
                    'etc': 'http://www.wps.cn/officeDocument/2017/etCustomData',
                    'xdr': 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                }
                for cell_img in root.findall('.//etc:cellImage', ns):
                    cnv_pr = cell_img.find('.//xdr:cNvPr', ns)
                    blip = cell_img.find('.//a:blip', ns)
                    if cnv_pr is not None and blip is not None:
                        img_name = cnv_pr.get('name', '')
                        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', '')
                        if img_name and rid:
                            name_to_rid[img_name] = rid

            rid_to_path = {}
            if 'xl/_rels/cellimages.xml.rels' in zf.namelist():
                rels_xml = zf.read('xl/_rels/cellimages.xml.rels')
                rels_root = ET.fromstring(rels_xml)
                for rel in rels_root:
                    rid = rel.get('Id', '')
                    target = rel.get('Target', '')
                    if rid and target:
                        rid_to_path[rid] = 'xl/' + target

            for img_name, rid in name_to_rid.items():
                fpath = rid_to_path.get(rid)
                if fpath and fpath in zf.namelist():
                    name_to_zip_path[img_name] = fpath

            zf.close()
        except Exception:
            pass

        # --- Load workbook with read_only=True (streams rows, low memory) ---
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        ws = wb.active

        # Read headers from first row
        for row in ws.iter_rows(min_row=1, max_row=1):
            header_row = [str(cell.value).strip() if cell.value else '' for cell in row]
            break

        # Count total rows for progress
        for _ in ws.iter_rows(min_row=2):
            total_rows += 1
        wb.close()

    col_mapping = {}  # col_idx (0-based) -> db_field
    screenshot_cols = {}  # col_idx (0-based) -> 'text' or 'ai'
    unknown_headers = []

    dynamic_mapping = get_dynamic_header_mapping(test_set_id)
    image_fields = {c["field"] for c in load_custom_columns(test_set_id) if c.get("type") == "image"}
    builtin_image_fields = {"text_screenshot_url", "ai_screenshot_url", "expected_screenshot"}
    for col_idx, header in enumerate(header_row):
        if not header:
            continue
        field = match_header(header, dynamic_mapping)
        if field == _IGNORE:
            continue
        if field:
            if field in builtin_image_fields or field in image_fields:
                # 图片列：既尝试内嵌图片(DISPIMG)提取，也保留纯 URL 文本路径
                screenshot_cols[col_idx] = field
                col_mapping[col_idx] = field
            else:
                col_mapping[col_idx] = field
        else:
            unknown_headers.append(header)

    if unknown_headers:
        if not skip_unknown:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return {"needs_confirmation": True, "unknown_columns": unknown_headers, "imported": 0, "duplicates_count": 0}

    # Find query column
    query_col = None
    for cidx, field in col_mapping.items():
        if field == 'query':
            query_col = cidx
            break
    if query_col is None:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return {"error": "未找到 query 列", "imported": 0, "duplicates_count": 0}

    import_tasks[task_id]["total"] = total_rows
    import_tasks[task_id]["status"] = "importing"

    # Start background import task
    if is_csv:
        asyncio.create_task(_do_import_csv(
            task_id, tmp_path, test_set_id, col_mapping, query_col, total_rows
        ))
    else:
        asyncio.create_task(_do_import(
            task_id, tmp_path, test_set_id, col_mapping, screenshot_cols,
            query_col, name_to_zip_path, total_rows
        ))

    return {"task_id": task_id, "total": total_rows, "status": "importing"}


async def _do_import(task_id, tmp_path, test_set_id, col_mapping, screenshot_cols, query_col, name_to_zip_path, total_rows):
    import re
    import zipfile

    try:
        # Load formula rows if needed
        formula_rows = []
        if screenshot_cols and name_to_zip_path:
            wb_formula = openpyxl.load_workbook(tmp_path, data_only=False, read_only=True)
            ws_formula = wb_formula.active
            for row in ws_formula.iter_rows(min_row=2):
                formula_row = {}
                for cidx in screenshot_cols:
                    cell = row[cidx] if cidx < len(row) else None
                    if cell and cell.value:
                        formula_row[cidx] = str(cell.value)
                formula_rows.append(formula_row)
            wb_formula.close()

        # Open data workbook for row iteration
        wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
        ws = wb.active

        db = await get_db()
        cursor = await db.execute("SELECT query FROM queries WHERE test_set_id = ?", [test_set_id])
        existing = {row[0] for row in await cursor.fetchall()}

        imported = 0
        duplicates = []
        now = datetime.now().isoformat()
        dispimg_re = re.compile(r'DISPIMG\("([^"]+)"')
        custom_cols = load_custom_columns(test_set_id)
        all_fields = EXCEL_COLUMNS + [c["field"] for c in custom_cols]
        field_names = ", ".join(all_fields + ["test_set_id", "created_at", "updated_at"])
        placeholders = ", ".join(["?"] * (len(all_fields) + 3))

        row_num = 0
        for row in ws.iter_rows(min_row=2):
            cells = [cell.value for cell in row]
            query_val = cells[query_col] if query_col < len(cells) else None
            if query_val is None or str(query_val).strip() == "":
                row_num += 1
                import_tasks[task_id]["progress"] = row_num
                continue

            query_str = str(query_val).strip()
            if query_str in existing:
                duplicates.append(query_str)
                row_num += 1
                import_tasks[task_id]["progress"] = row_num
                continue

            row_dict = {f: None for f in EXCEL_COLUMNS}
            for col_idx, field in col_mapping.items():
                val = cells[col_idx] if col_idx < len(cells) else None
                if isinstance(val, str) and val.startswith("="):
                    val = None
                row_dict[field] = val

            row_data_full = [row_dict.get(f) for f in all_fields]

            await db.execute(
                f"INSERT INTO queries ({field_names}) VALUES ({placeholders})",
                row_data_full + [test_set_id, now, now]
            )
            existing.add(query_str)
            imported += 1

            new_id_cursor = await db.execute("SELECT last_insert_rowid()")
            new_id = (await new_id_cursor.fetchone())[0]

            # Extract DISPIMG images on-demand from ZIP
            if formula_rows and row_num < len(formula_rows):
                formula_row = formula_rows[row_num]
                for col_idx, target_field in screenshot_cols.items():
                    formula_val = formula_row.get(col_idx, '')
                    if 'DISPIMG' in formula_val:
                        m = dispimg_re.search(formula_val)
                        if m:
                            img_id = m.group(1)
                            zip_path = name_to_zip_path.get(img_id)
                            if zip_path:
                                zf = zipfile.ZipFile(tmp_path)
                                img_bytes = zf.read(zip_path)
                                zf.close()
                                ext = 'png' if img_bytes[:4] == b'\x89PNG' else 'jpg'
                                suffix = re.sub(r'[^0-9A-Za-z_]', '', target_field) or 'img'
                                filename = f"q{new_id}_{suffix}_{int(datetime.now().timestamp())}.{ext}"
                                filepath = os.path.join(SCREENSHOTS_DIR, filename)
                                with open(filepath, "wb") as f:
                                    f.write(img_bytes)
                                url = f"/static/screenshots/{filename}"
                                await db.execute(f'UPDATE queries SET "{target_field}" = ? WHERE id = ?', [url, new_id])

            row_num += 1
            import_tasks[task_id]["progress"] = row_num
            import_tasks[task_id]["imported"] = imported

            # Commit every 100 rows
            if imported % 100 == 0:
                await db.commit()
                await asyncio.sleep(0)  # yield to event loop

        await db.commit()
        wb.close()
        await db.close()

        result = {"imported": imported, "duplicates_count": len(duplicates)}
        if duplicates:
            result["duplicates"] = duplicates

        import_tasks[task_id]["status"] = "done"
        import_tasks[task_id]["result"] = result
        await manager.broadcast({"type": "data_refresh"})

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[IMPORT ERROR] {e}", flush=True)
        import_tasks[task_id]["status"] = "error"
        import_tasks[task_id]["error"] = str(e)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _read_csv_rows(path):
    """读取 CSV 全部行，自动尝试常见中文编码。返回 list[list[str]]。"""
    for enc in ("utf-8-sig", "gbk", "gb18030", "utf-8"):
        try:
            with open(path, "r", encoding=enc, newline="") as f:
                return list(csv.reader(f))
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        return list(csv.reader(f))


async def _do_import_csv(task_id, tmp_path, test_set_id, col_mapping, query_col, total_rows):
    try:
        csv_rows = _read_csv_rows(tmp_path)
        data_rows = csv_rows[1:] if csv_rows else []

        db = await get_db()
        cursor = await db.execute("SELECT query FROM queries WHERE test_set_id = ?", [test_set_id])
        existing = {row[0] for row in await cursor.fetchall()}

        imported = 0
        duplicates = []
        now = datetime.now().isoformat()
        custom_cols = load_custom_columns(test_set_id)
        all_fields = EXCEL_COLUMNS + [c["field"] for c in custom_cols]
        field_names = ", ".join(all_fields + ["test_set_id", "created_at", "updated_at"])
        placeholders = ", ".join(["?"] * (len(all_fields) + 3))

        row_num = 0
        for cells in data_rows:
            query_val = cells[query_col] if query_col < len(cells) else None
            if query_val is None or str(query_val).strip() == "":
                row_num += 1
                import_tasks[task_id]["progress"] = row_num
                continue

            query_str = str(query_val).strip()
            if query_str in existing:
                duplicates.append(query_str)
                row_num += 1
                import_tasks[task_id]["progress"] = row_num
                continue

            row_dict = {f: None for f in EXCEL_COLUMNS}
            for col_idx, field in col_mapping.items():
                val = cells[col_idx] if col_idx < len(cells) else None
                if isinstance(val, str):
                    val = val.strip()
                    if val == "" or val.startswith("="):
                        val = None
                row_dict[field] = val

            row_data_full = [row_dict.get(f) for f in all_fields]
            await db.execute(
                f"INSERT INTO queries ({field_names}) VALUES ({placeholders})",
                row_data_full + [test_set_id, now, now]
            )
            existing.add(query_str)
            imported += 1

            row_num += 1
            import_tasks[task_id]["progress"] = row_num
            import_tasks[task_id]["imported"] = imported

            if imported % 100 == 0:
                await db.commit()
                await asyncio.sleep(0)

        await db.commit()
        await db.close()

        result = {"imported": imported, "duplicates_count": len(duplicates)}
        if duplicates:
            result["duplicates"] = duplicates

        import_tasks[task_id]["status"] = "done"
        import_tasks[task_id]["result"] = result
        await manager.broadcast({"type": "data_refresh"})

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[CSV IMPORT ERROR] {e}", flush=True)
        import_tasks[task_id]["status"] = "error"
        import_tasks[task_id]["error"] = str(e)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


@app.get("/api/import/progress")
async def get_import_progress(task_id: str):
    task = import_tasks.get(task_id)
    if not task:
        return {"error": "task not found"}
    return task


@app.post("/api/import/local")
async def import_local_file(test_set_id: int = 1, path: str = "/tmp/import.xlsx", skip_unknown: int = 1):
    """Import from a file already on the server (uploaded via SCP)."""
    if not os.path.exists(path):
        return {"error": f"文件不存在: {path}"}

    import re
    import zipfile
    from xml.etree import ElementTree as ET

    task_id = str(uuid.uuid4())[:8]
    import_tasks[task_id] = {"status": "parsing", "progress": 0, "total": 0, "imported": 0, "error": None, "result": None}
    tmp_path = path

    # --- Extract DISPIMG name->zip_path mapping ---
    name_to_zip_path = {}
    try:
        zf = zipfile.ZipFile(tmp_path)
        name_to_rid = {}
        if 'xl/cellimages.xml' in zf.namelist():
            ci_xml = zf.read('xl/cellimages.xml')
            root = ET.fromstring(ci_xml)
            ns = {
                'etc': 'http://www.wps.cn/officeDocument/2017/etCustomData',
                'xdr': 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            }
            for cell_img in root.findall('.//etc:cellImage', ns):
                cnv_pr = cell_img.find('.//xdr:cNvPr', ns)
                blip = cell_img.find('.//a:blip', ns)
                if cnv_pr is not None and blip is not None:
                    img_name = cnv_pr.get('name', '')
                    rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', '')
                    if img_name and rid:
                        name_to_rid[img_name] = rid

        rid_to_path = {}
        if 'xl/_rels/cellimages.xml.rels' in zf.namelist():
            rels_xml = zf.read('xl/_rels/cellimages.xml.rels')
            rels_root = ET.fromstring(rels_xml)
            for rel in rels_root:
                rid = rel.get('Id', '')
                target = rel.get('Target', '')
                if rid and target:
                    rid_to_path[rid] = 'xl/' + target

        for img_name, rid in name_to_rid.items():
            fpath = rid_to_path.get(rid)
            if fpath and fpath in zf.namelist():
                name_to_zip_path[img_name] = fpath
        zf.close()
    except Exception:
        pass

    wb = openpyxl.load_workbook(tmp_path, data_only=True, read_only=True)
    ws = wb.active

    header_row = []
    for row in ws.iter_rows(min_row=1, max_row=1):
        header_row = [str(cell.value).strip() if cell.value else '' for cell in row]
        break

    col_mapping = {}
    screenshot_cols = {}
    dynamic_mapping = get_dynamic_header_mapping(test_set_id)
    for col_idx, header in enumerate(header_row):
        if not header:
            continue
        field = dynamic_mapping.get(header)
        if field:
            if field == 'text_screenshot_url':
                screenshot_cols[col_idx] = 'text'
            elif field == 'ai_screenshot_url':
                screenshot_cols[col_idx] = 'ai'
            else:
                col_mapping[col_idx] = field
        else:
            for k, v in dynamic_mapping.items():
                if v and (k.lower() in header.lower() or header.lower() in k.lower()):
                    col_mapping[col_idx] = v
                    break

    total_rows = 0
    for _ in ws.iter_rows(min_row=2):
        total_rows += 1
    wb.close()

    query_col = None
    for cidx, field in col_mapping.items():
        if field == 'query':
            query_col = cidx
            break
    if query_col is None:
        return {"error": "未找到 query 列"}

    import_tasks[task_id]["total"] = total_rows
    import_tasks[task_id]["status"] = "importing"

    asyncio.create_task(_do_import(
        task_id, tmp_path, test_set_id, col_mapping, screenshot_cols,
        query_col, name_to_zip_path, total_rows
    ))

    return {"task_id": task_id, "total": total_rows, "status": "importing"}


@app.get("/api/export/excel")
async def export_excel():
    from openpyxl.drawing.image import Image as XlImage
    from openpyxl.utils import get_column_letter
    from PIL import Image as PILImage

    db = await get_db()
    cursor = await db.execute("SELECT * FROM queries ORDER BY id")
    rows = await cursor.fetchall()
    await db.close()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    ws.append(EXCEL_HEADERS)

    text_col_idx = EXCEL_COLUMNS.index("text_screenshot_url") + 1
    ai_col_idx = EXCEL_COLUMNS.index("ai_screenshot_url") + 1

    for row_idx, row in enumerate(rows, start=2):
        row_dict = dict(row)
        row_data = []
        for col in EXCEL_COLUMNS:
            val = row_dict.get(col)
            if col in ("text_screenshot_url", "ai_screenshot_url"):
                row_data.append(None)
            else:
                row_data.append(val)
        ws.append(row_data)

        for col_idx, field in [(text_col_idx, "text_screenshot_url"), (ai_col_idx, "ai_screenshot_url")]:
            url = row_dict.get(field)
            if url:
                img_path = os.path.join(APP_DIR, url.lstrip("/").replace("/", os.sep))
                if os.path.exists(img_path):
                    try:
                        pil_img = PILImage.open(img_path)
                        if pil_img.width > 960:
                            ratio = 960 / pil_img.width
                            pil_img = pil_img.resize((960, int(pil_img.height * ratio)), PILImage.LANCZOS)
                        pil_img = pil_img.convert("RGB")
                        tmp_buf = BytesIO()
                        pil_img.save(tmp_buf, "JPEG", quality=80)
                        tmp_buf.seek(0)

                        img = XlImage(tmp_buf)
                        cell_w = 180
                        img.height = int(cell_w * pil_img.height / pil_img.width)
                        img.width = cell_w
                        cell_ref = f"{get_column_letter(col_idx)}{row_idx}"
                        ws.add_image(img, cell_ref)
                        ws.row_dimensions[row_idx].height = max(ws.row_dimensions[row_idx].height or 15, img.height * 0.75)
                    except Exception:
                        pass

    ws.column_dimensions[get_column_letter(text_col_idx)].width = 28
    ws.column_dimensions[get_column_letter(ai_col_idx)].width = 28

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=export.xlsx"}
    )


@app.get("/api/export/csv")
async def export_csv():
    db = await get_db()
    cursor = await db.execute("SELECT * FROM queries ORDER BY id")
    rows = await cursor.fetchall()
    await db.close()

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(EXCEL_HEADERS)

    for row in rows:
        row_dict = dict(row)
        writer.writerow([row_dict.get(col) for col in EXCEL_COLUMNS])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=export.csv"}
    )


@app.delete("/api/queries/{query_id}")
async def delete_query(query_id: int):
    db = await get_db()
    await db.execute("DELETE FROM queries WHERE id = ?", [query_id])
    await db.commit()
    await db.close()
    return {"deleted": query_id}


@app.delete("/api/queries")
async def delete_all_queries(no_backup: int = 0, test_set_id: int = None):
    import shutil
    import glob

    backup_name = None
    if not no_backup:
        backup_dir = os.path.join(APP_DIR, "backups")
        os.makedirs(backup_dir, exist_ok=True)
        backup_name = f"data_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
        shutil.copy2(DB_PATH, os.path.join(backup_dir, backup_name))

        for f in glob.glob(os.path.join(backup_dir, "data_backup_*.db")):
            age = time.time() - os.path.getmtime(f)
            if age > 7 * 86400:
                os.remove(f)

    db = await get_db()
    if test_set_id:
        cursor = await db.execute("SELECT COUNT(*) FROM queries WHERE test_set_id = ?", [test_set_id])
        count = (await cursor.fetchone())[0]
        await db.execute("DELETE FROM queries WHERE test_set_id = ?", [test_set_id])
    else:
        cursor = await db.execute("SELECT COUNT(*) FROM queries")
        count = (await cursor.fetchone())[0]
        await db.execute("DELETE FROM queries")
    try:
        await db.execute("DELETE FROM sqlite_sequence WHERE name='queries'")
    except Exception:
        pass
    await db.commit()
    await db.close()
    await manager.broadcast({"type": "data_refresh"})
    result = {"deleted": count}
    if backup_name:
        result["backup"] = backup_name
    return result


@app.post("/api/restore")
async def restore_backup(filename: str):
    """从备份恢复数据"""
    import shutil
    backup_path = os.path.join(APP_DIR, "backups", filename)
    if not os.path.exists(backup_path):
        raise HTTPException(404, "Backup not found")
    shutil.copy2(backup_path, DB_PATH)
    await manager.broadcast({"type": "data_refresh"})
    return {"restored": filename}


@app.get("/api/backups")
async def list_backups():
    """列出所有可用备份"""
    backup_dir = os.path.join(APP_DIR, "backups")
    if not os.path.exists(backup_dir):
        return []
    import glob
    files = glob.glob(os.path.join(backup_dir, "data_backup_*.db"))
    files.sort(key=os.path.getmtime, reverse=True)
    return [{"filename": os.path.basename(f), "time": datetime.fromtimestamp(os.path.getmtime(f)).strftime('%Y-%m-%d %H:%M:%S')} for f in files]


@app.post("/api/queries")
async def create_query(data: QueryUpdate):
    fields = data.model_dump(exclude_unset=True)
    now = datetime.now().isoformat()
    fields["created_at"] = now
    fields["updated_at"] = now

    cols = ", ".join(fields.keys())
    placeholders = ", ".join("?" * len(fields))
    values = list(fields.values())

    db = await get_db()
    cursor = await db.execute(f"INSERT INTO queries ({cols}) VALUES ({placeholders})", values)
    await db.commit()
    new_id = cursor.lastrowid

    cursor = await db.execute("SELECT * FROM queries WHERE id = ?", [new_id])
    row = await cursor.fetchone()
    await db.close()
    return dict(row)


# ============================================================
# Analysis API
# ============================================================

@app.get("/api/analysis")
async def get_analysis(test_set_id: int = QueryParam(default=None)):
    db = await get_db()
    if test_set_id:
        cursor = await db.execute("SELECT * FROM queries WHERE test_set_id = ?", [test_set_id])
    else:
        cursor = await db.execute("SELECT * FROM queries")
    rows = [dict(r) for r in await cursor.fetchall()]
    await db.close()

    import statistics

    score_fields = ['relevance', 'accuracy', 'ranking', 'diversity', 'quality_threshold', 'copy_recommendation', 'click_desire']

    def to_float(v):
        if v is None: return None
        try: return float(v)
        except (ValueError, TypeError): return None

    # Overview
    total = len(rows)
    scored = [r for r in rows if to_float(r.get('relevance')) is not None]
    scored_count = len(scored)
    ai_scores = [to_float(r['ai_search_score']) for r in rows if to_float(r.get('ai_search_score')) is not None]

    overview = {
        "total": total,
        "scored": scored_count,
        "ai_avg": round(statistics.mean(ai_scores), 2) if ai_scores else 0,
    }

    # Dimensions
    dimensions = []
    for f in score_fields:
        vals = [to_float(r[f]) for r in rows if to_float(r.get(f)) is not None]
        if vals:
            dimensions.append({
                "name": f,
                "mean": round(statistics.mean(vals), 2),
                "min": min(vals),
                "max": max(vals),
                "count": len(vals),
            })
        else:
            dimensions.append({"name": f, "mean": 0, "min": 0, "max": 0, "count": 0})

    # VS text search
    vs_vals = [r['vs_text_search'] for r in rows if r.get('vs_text_search') is not None]
    better = sum(1 for v in vs_vals if str(v).strip() == '1')
    same = sum(1 for v in vs_vals if str(v).strip() == '0')
    worse = sum(1 for v in vs_vals if str(v).strip() == '-1')
    vs_text = {"better": better, "same": same, "worse": worse, "total": better + same + worse}

    # By category
    from collections import defaultdict
    cat_data = defaultdict(list)
    for r in rows:
        cat = r.get('business_category') or '未分类'
        cat_data[cat].append(r)

    by_category = []
    for cat, cat_rows in sorted(cat_data.items(), key=lambda x: -len(x[1])):
        cat_info = {"category": cat, "n": len(cat_rows), "dims": {}}
        for f in score_fields:
            vals = [to_float(r[f]) for r in cat_rows if to_float(r.get(f)) is not None]
            cat_info["dims"][f] = round(statistics.mean(vals), 2) if vals else 0
        ai = [to_float(r['ai_search_score']) for r in cat_rows if to_float(r.get('ai_search_score')) is not None]
        cat_info["ai_avg"] = round(statistics.mean(ai), 2) if ai else 0
        # VS for this category
        cat_vs = [r['vs_text_search'] for r in cat_rows if r.get('vs_text_search') is not None]
        cat_info["vs"] = {
            "better": sum(1 for v in cat_vs if str(v).strip() == '1'),
            "same": sum(1 for v in cat_vs if str(v).strip() == '0'),
            "worse": sum(1 for v in cat_vs if str(v).strip() == '-1'),
        }
        by_category.append(cat_info)

    # By query level (short/long based on is_short_query field)
    short_rows = [r for r in rows if r.get('is_short_query') and '短' in str(r['is_short_query'])]
    long_rows = [r for r in rows if r.get('is_short_query') and '长' in str(r['is_short_query'])]

    def calc_group_dims(group_rows):
        result = {}
        for f in score_fields:
            vals = [to_float(r[f]) for r in group_rows if to_float(r.get(f)) is not None]
            result[f] = round(statistics.mean(vals), 2) if vals else 0
        ai = [to_float(r['ai_search_score']) for r in group_rows if to_float(r.get('ai_search_score')) is not None]
        result['ai_avg'] = round(statistics.mean(ai), 2) if ai else 0
        return result

    by_query_level = {
        "short": {"n": len(short_rows), "dims": calc_group_dims(short_rows)},
        "long": {"n": len(long_rows), "dims": calc_group_dims(long_rows)},
    }

    # By frequency
    freq_groups = defaultdict(list)
    for r in rows:
        freq = r.get('frequency_level') or '未知'
        freq_groups[freq].append(r)

    by_frequency = {}
    for freq, freq_rows in freq_groups.items():
        by_frequency[freq] = {"n": len(freq_rows), "dims": calc_group_dims(freq_rows)}

    # TOP5 / BOTTOM5
    scored_with_ai = [r for r in rows if to_float(r.get('ai_search_score')) is not None]
    sorted_by_ai = sorted(scored_with_ai, key=lambda r: to_float(r['ai_search_score']), reverse=True)

    def pick_fields(r):
        return {"id": r["id"], "query": r["query"], "ai_score": to_float(r["ai_search_score"]),
                "category": r.get("business_category")}

    top5 = [pick_fields(r) for r in sorted_by_ai[:5]]
    bottom5 = [pick_fields(r) for r in sorted_by_ai[-5:]]

    return {
        "overview": overview,
        "dimensions": dimensions,
        "vs_text": vs_text,
        "by_category": by_category,
        "by_query_level": by_query_level,
        "by_frequency": by_frequency,
        "top5": top5,
        "bottom5": bottom5,
    }


@app.post("/api/analysis/ai-text")
async def ai_text_analysis():
    """Call Kimi API for deep analysis"""
    import httpx

    api_key = os.environ.get("KIMI_API_KEY", "")
    if not api_key:
        raise HTTPException(400, "KIMI_API_KEY not configured")

    # Get analysis data
    from starlette.testclient import TestClient
    # Simpler: just call get_analysis directly
    analysis = await get_analysis()

    prompt = f"""你是一个搜索质量评测专家。以下是 AI 搜索评测的统计数据，请给出专业的分析报告（中文），包括：
1. 总体评价
2. 各维度优劣势分析
3. 与文本搜索对比的结论
4. 按业务分类的差异化表现
5. 改进建议

数据：
- 总览：{json.dumps(analysis['overview'], ensure_ascii=False)}
- 各维度均分：{json.dumps(analysis['dimensions'], ensure_ascii=False)}
- AI搜 vs 文本搜：{json.dumps(analysis['vs_text'], ensure_ascii=False)}
- 按业务分类：{json.dumps(analysis['by_category'][:6], ensure_ascii=False)}
"""

    async with httpx.AsyncClient(timeout=60) as client:
        resp = await client.post(
            "https://api.moonshot.cn/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}"},
            json={
                "model": "moonshot-v1-8k",
                "messages": [{"role": "user", "content": prompt}],
            }
        )
        resp.raise_for_status()
        data = resp.json()
        return {"text": data["choices"][0]["message"]["content"]}


SCORE_FIELDS = ['relevance', 'ranking', 'diversity', 'quality_threshold']
DIM_LABELS = {
    'relevance': '相关性',
    'ranking': '排序合理性',
    'diversity': '多样性',
    'quality_threshold': '质量门槛（与文本搜索对比）',
}


def build_report_model(rows):
    """计算文档2格式报告所需的全部结构化数据（供网页展示与docx共用）"""
    def to_float(v):
        if v is None:
            return None
        try:
            return float(v)
        except (ValueError, TypeError):
            return None

    total = len(rows)

    # AI综合分 = 4维度均值（逐条）
    for r in rows:
        vals = [to_float(r.get(f)) for f in SCORE_FIELDS]
        vals = [v for v in vals if v is not None]
        r['_ai_avg'] = sum(vals) / len(vals) if vals else None

    # 维度统计
    dim_stats = {}
    for f in SCORE_FIELDS:
        vals = [to_float(r.get(f)) for r in rows]
        vals = [v for v in vals if v is not None]
        if vals:
            dim_stats[f] = {
                'mean': sum(vals) / len(vals),
                'count': len(vals),
                'low_rate': sum(1 for v in vals if v <= 2) / len(vals),
            }

    def vs_counts(sub):
        w = sum(1 for r in sub if str(r.get('vs_text_search')).strip() == '1')
        d = sum(1 for r in sub if str(r.get('vs_text_search')).strip() == '0')
        l = sum(1 for r in sub if str(r.get('vs_text_search')).strip() == '-1')
        return w, d, l

    total_win, total_draw, total_lose = vs_counts(rows)
    total_vs = total_win + total_draw + total_lose
    win_rate = total_win / total_vs * 100 if total_vs else 0
    lose_rate = total_lose / total_vs * 100 if total_vs else 0

    from collections import defaultdict
    cat_groups = defaultdict(list)
    for r in rows:
        cat_groups[r.get('business_category') or '未分类'].append(r)

    cat_data = {}
    for cat, sub in cat_groups.items():
        w, d, l = vs_counts(sub)
        scores = {}
        for f in SCORE_FIELDS:
            vals = [to_float(x.get(f)) for x in sub]
            vals = [v for v in vals if v is not None]
            scores[f] = sum(vals) / len(vals) if vals else 0
        cat_data[cat] = {'n': len(sub), 'win': w, 'draw': d, 'lose': l,
                         'vs_n': w + d + l, 'scores': scores}
    categories = sorted(cat_data.keys(), key=lambda c: -cat_data[c]['n'])

    # 强/弱分类
    strong = [(c, cat_data[c]) for c in categories
              if cat_data[c]['vs_n'] >= 3 and cat_data[c]['win'] > cat_data[c]['lose']]
    weak = [(c, cat_data[c]) for c in categories
            if cat_data[c]['vs_n'] >= 3 and cat_data[c]['lose'] > cat_data[c]['win']]
    worst = max(dim_stats.items(), key=lambda kv: kv[1]['low_rate']) if dim_stats else None

    # 特征切分
    def feat_stats(sub):
        n = len(sub)
        ai_vals = [r['_ai_avg'] for r in sub if r['_ai_avg'] is not None]
        ai = sum(ai_vals) / len(ai_vals) if ai_vals else None
        out = {'n': n, 'ai': ai}
        for f in ['relevance', 'ranking', 'quality_threshold']:
            vs = [to_float(r.get(f)) for r in sub]
            vs = [v for v in vs if v is not None]
            out[f] = sum(vs) / len(vs) if vs else None
        return out

    features = []
    short_rows = [r for r in rows if r.get('is_short_query') and '短' in str(r['is_short_query'])]
    long_rows = [r for r in rows if r.get('is_short_query') and '长' in str(r['is_short_query'])]
    features.append({'dim': '长短', 'group': '短query', **feat_stats(short_rows)})
    features.append({'dim': '长短', 'group': '长query', **feat_stats(long_rows)})
    freq_groups = defaultdict(list)
    for r in rows:
        if r.get('frequency_level'):
            freq_groups[str(r['frequency_level'])].append(r)
    for g in ['高频', '中频', '长尾']:
        if g in freq_groups:
            features.append({'dim': '频率', 'group': g, **feat_stats(freq_groups[g])})

    return {
        'total': total,
        'n_categories': len(categories),
        'conclusion': {
            'trend': '优于' if win_rate >= lose_rate else '弱于',
            'win_rate': round(win_rate),
            'lose_rate': round(lose_rate),
            'worst': ({'label': DIM_LABELS[worst[0]].split('（')[0],
                       'mean': round(worst[1]['mean'], 2),
                       'low_rate': round(worst[1]['low_rate'] * 100)} if worst else None),
        },
        'strong': [{'cat': c, 'win': d['win'], 'vs_n': d['vs_n']} for c, d in strong[:6]],
        'weak': [{'cat': c, 'win': d['win'], 'vs_n': d['vs_n']} for c, d in weak[:6]],
        'overall_vs': {
            'win': total_win, 'draw': total_draw, 'lose': total_lose, 'total': total_vs,
            'win_rate': round(win_rate),
            'draw_rate': round(total_draw / total_vs * 100) if total_vs else 0,
            'lose_rate': round(lose_rate),
        },
        'dimensions': [
            {'field': f, 'label': DIM_LABELS[f],
             'mean': round(st['mean'], 2), 'count': st['count'],
             'low_rate': round(st['low_rate'] * 100),
             'note': ('⚠️ 重点改进' if st['low_rate'] >= 0.4 else ('✓ 较好' if st['low_rate'] <= 0.15 else ''))}
            for f, st in sorted(dim_stats.items(), key=lambda kv: kv[1]['mean'])
        ],
        'features': [
            {'dim': ft['dim'], 'group': ft['group'], 'n': ft['n'],
             'ai': round(ft['ai'], 2) if ft['ai'] is not None else None,
             'relevance': round(ft['relevance'], 2) if ft['relevance'] is not None else None,
             'ranking': round(ft['ranking'], 2) if ft['ranking'] is not None else None,
             'quality': round(ft['quality_threshold'], 2) if ft['quality_threshold'] is not None else None}
            for ft in features
        ],
        'categories': [
            {'cat': c, 'n': cat_data[c]['n'],
             'relevance': round(cat_data[c]['scores'].get('relevance', 0), 2),
             'ranking': round(cat_data[c]['scores'].get('ranking', 0), 2),
             'quality': round(cat_data[c]['scores'].get('quality_threshold', 0), 2),
             'win': cat_data[c]['win'], 'draw': cat_data[c]['draw'],
             'lose': cat_data[c]['lose'], 'vs_n': cat_data[c]['vs_n']}
            for c in categories
        ],
    }


async def _load_report_rows(test_set_id):
    db = await get_db()
    if test_set_id:
        cursor = await db.execute("SELECT * FROM queries WHERE test_set_id = ?", [test_set_id])
    else:
        cursor = await db.execute("SELECT * FROM queries")
    rows = [dict(r) for r in await cursor.fetchall()]
    await db.close()
    return rows


async def _prev_dim_means(test_set_id):
    """Mean per score dimension for the previous test set (id just below)."""
    if not test_set_id:
        return None
    db = await get_db()
    cur = await db.execute("SELECT id FROM test_sets WHERE id < ? ORDER BY id DESC LIMIT 1", [test_set_id])
    row = await cur.fetchone()
    if not row:
        await db.close()
        return None
    prev_id = row[0]
    cur = await db.execute("SELECT * FROM queries WHERE test_set_id = ?", [prev_id])
    prows = [dict(r) for r in await cur.fetchall()]
    await db.close()
    means = {}
    for f in SCORE_FIELDS:
        vals = [r[f] for r in prows if r.get(f) is not None]
        if vals:
            means[f] = sum(vals) / len(vals)
    return means


@app.get("/api/report-data")
async def report_data(test_set_id: int = QueryParam(default=None)):
    """返回结构化报告数据，供网页直接展示"""
    rows = await _load_report_rows(test_set_id)
    model = build_report_model(rows)
    prev_means = await _prev_dim_means(test_set_id)
    if prev_means:
        for d in model['dimensions']:
            pm = prev_means.get(d['field'])
            d['delta'] = round(d['mean'] - pm, 2) if pm is not None else None
    model['notes'] = load_report_notes(test_set_id)
    return model


class ReportNotes(BaseModel):
    conclusion: str = ""
    priorities: str = ""
    stats: str = ""
    badcase: str = ""


@app.get("/api/report-notes")
async def get_report_notes(test_set_id: int = QueryParam(default=None)):
    return load_report_notes(test_set_id)


@app.post("/api/report-notes")
async def post_report_notes(notes: ReportNotes, test_set_id: int = QueryParam(default=None)):
    return save_report_notes(notes.dict(), test_set_id)


def render_badcase_html(doc, html_str):
    """Render bad-case rich-text HTML (tables + paragraphs) into a docx document."""
    import re as _re
    import html as _htmllib
    from html.parser import HTMLParser

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self.blocks = []
            self.in_table = False
            self.table = None
            self.row = None
            self.cell = None
            self.para = []

        def _flush(self):
            t = ''.join(self.para).strip()
            if t:
                self.blocks.append(('para', t))
            self.para = []

        def handle_starttag(self, tag, attrs):
            if tag == 'table':
                self._flush(); self.in_table = True; self.table = []
            elif tag == 'tr' and self.in_table:
                self.row = []
            elif tag in ('td', 'th') and self.in_table:
                self.cell = []
            elif tag == 'br':
                (self.cell if (self.in_table and self.cell is not None) else self.para).append('\n')

        def handle_endtag(self, tag):
            if tag == 'table' and self.in_table:
                self.blocks.append(('table', self.table)); self.in_table = False; self.table = None
            elif tag == 'tr' and self.in_table and self.row is not None:
                self.table.append(self.row); self.row = None
            elif tag in ('td', 'th') and self.in_table and self.cell is not None:
                if self.row is not None:
                    self.row.append(''.join(self.cell).strip())
                self.cell = None
            elif tag in ('p', 'div') and not self.in_table:
                self._flush()

        def handle_data(self, data):
            if self.in_table and self.cell is not None:
                self.cell.append(data)
            elif not self.in_table:
                self.para.append(data)

    p = _P()
    try:
        p.feed(html_str or '')
    except Exception:
        pass
    p._flush()

    if not p.blocks:
        text = _htmllib.unescape(_re.sub(r'<[^>]+>', '', html_str or '')).strip()
        if text:
            doc.add_paragraph(text)
        return

    for kind, payload in p.blocks:
        if kind == 'para':
            doc.add_paragraph(_htmllib.unescape(payload))
        elif kind == 'table' and payload:
            ncols = max((len(r) for r in payload), default=0)
            if ncols == 0:
                continue
            t = doc.add_table(rows=0, cols=ncols)
            t.style = 'Table Grid'
            for r in payload:
                cells = t.add_row().cells
                for i in range(ncols):
                    cells[i].text = _htmllib.unescape(r[i]) if i < len(r) else ''


@app.post("/api/export/report")
async def export_report(test_set_id: int = QueryParam(default=None)):
    """生成文档2格式的Word评测报告（图表+特征切分+bad case框架）"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    matplotlib.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
    matplotlib.rcParams['axes.unicode_minus'] = False
    import numpy as np

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    rows = await _load_report_rows(test_set_id)
    model = build_report_model(rows)
    notes = load_report_notes(test_set_id)
    total = model['total']
    categories = model['categories']
    total_vs = model['overall_vs']['total']

    # cat_data 便于绘图
    cat_data = {c['cat']: c for c in categories}
    cat_names = [c['cat'] for c in categories]

    # --- Chart 1: 各业务类型关键维度得分 ---
    top_cats = [c for c in cat_names if cat_data[c]['n'] >= 5][:6]
    plot_dims = [('relevance', '相关性'), ('ranking', '排序合理性'), ('quality', '质量门槛')]
    colors = ['#5470c6', '#91cc75', '#fac858', '#ee6666']

    chart1_buf = None
    if top_cats:
        fig, ax = plt.subplots(figsize=(10, 5))
        x = np.arange(len(top_cats))
        width = 0.8 / len(plot_dims)
        for idx, (key, label) in enumerate(plot_dims):
            vals = [cat_data[c].get(key, 0) for c in top_cats]
            bars = ax.bar(x + idx * width - (len(plot_dims) - 1) * width / 2, vals, width,
                          label=label, color=colors[idx % len(colors)])
            for bar in bars:
                h = bar.get_height()
                if h > 0:
                    ax.annotate(f'{h:.1f}', xy=(bar.get_x() + bar.get_width() / 2, h),
                                xytext=(0, 3), textcoords="offset points",
                                ha='center', va='bottom', fontsize=8)
        ax.set_ylabel('得分 (满分5分)')
        ax.set_title('各业务类型关键维度得分', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels([f'{c}\n(n={cat_data[c]["n"]})' for c in top_cats])
        ax.set_ylim(0, 5.5)
        ax.axhline(y=3, color='red', linestyle='--', linewidth=0.8, alpha=0.5, label='及格线(3分)')
        ax.legend(loc='upper right')
        plt.tight_layout()
        chart1_buf = BytesIO()
        plt.savefig(chart1_buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        chart1_buf.seek(0)

    # --- Chart 2: vs文本搜 净胜率（横向堆叠） ---
    plot_cats = [c for c in cat_names if cat_data[c]['vs_n'] >= 3][:8]
    plot_cats.reverse()
    chart2_buf = None
    if plot_cats:
        wins = [cat_data[c]['win'] for c in plot_cats]
        draws = [cat_data[c]['draw'] for c in plot_cats]
        loses = [cat_data[c]['lose'] for c in plot_cats]
        fig2, ax2 = plt.subplots(figsize=(8, 4))
        y = np.arange(len(plot_cats))
        bh = 0.6
        ax2.barh(y, wins, bh, label='AI优', color='#91cc75')
        ax2.barh(y, draws, bh, left=wins, label='持平', color='#fac858')
        ax2.barh(y, loses, bh, left=[w + d for w, d in zip(wins, draws)], label='AI劣', color='#ee6666')
        ax2.set_yticks(y)
        ax2.set_yticklabels(plot_cats)
        ax2.set_xlabel('query数量')
        ax2.set_title('AI搜索 vs 文本搜索 对比结果', fontsize=13, fontweight='bold')
        ax2.legend(loc='lower right')
        for i in range(len(plot_cats)):
            t = wins[i] + draws[i] + loses[i]
            if t > 0:
                net = (wins[i] - loses[i]) / t * 100
                if net > 0:
                    label, color = f'+{net:.0f}%', '#2e7d32'
                elif net < 0:
                    label, color = f'{net:.0f}%', '#c62828'
                else:
                    label, color = '±0%', '#666666'
                ax2.text(t + 0.3, i, label, va='center', fontsize=9, color=color, fontweight='bold')
        plt.tight_layout()
        chart2_buf = BytesIO()
        plt.savefig(chart2_buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        chart2_buf.seek(0)

    # ========================================================
    # WORD
    # ========================================================
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    title = doc.add_heading('AI搜索评测报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f'评测样本：{total}条query（覆盖{model["n_categories"]}种业务类型） | '
        f'评分：4维度×5分制 | 基准对比：文本搜索'
    )

    # === 一、核心结论 ===
    doc.add_heading('一、核心结论', level=1)
    concl_m = model['conclusion']
    ovs = model['overall_vs']
    p = doc.add_paragraph()
    concl = f'AI搜索整体{concl_m["trend"]}文本搜索（胜率{concl_m["win_rate"]}% vs {concl_m["lose_rate"]}%）'
    if concl_m['worst']:
        w = concl_m['worst']
        concl += f'，{w["label"]}维度表现不足（均分{w["mean"]}，{w["low_rate"]}%低分率）。'
    else:
        concl += '。'
    p.add_run(concl).bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('按query业务分类：').bold = True
    if model['strong']:
        doc.add_paragraph('AI搜索强于文本搜：' + '、'.join(f'{s["cat"]}(胜率{s["win"]}/{s["vs_n"]})' for s in model['strong'][:4]))
    if model['weak']:
        doc.add_paragraph('AI搜索弱于文本搜：' + '、'.join(f'{s["cat"]}(胜率{s["win"]}/{s["vs_n"]})' for s in model['weak'][:4]))

    if notes.get('conclusion', '').strip():
        for line in notes['conclusion'].strip().split('\n'):
            doc.add_paragraph(line)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('本周可重点提升：').bold = True
    if notes.get('priorities', '').strip():
        for line in notes['priorities'].strip().split('\n'):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph('（此处根据bad case人工总结，例：地图质量 / 相关性召回 / 缺图漏图）')

    # === 二、评测数据总览 ===
    doc.add_heading('二、评测数据总览', level=1)

    # 2.1 统计数据
    p = doc.add_paragraph()
    p.add_run('2.1 统计数据').bold = True
    if notes.get('stats', '').strip():
        for line in notes['stats'].strip().split('\n'):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph('（点击率QV、游玩平均时长、复玩率等线上业务指标图表，请手动补充）')

    # 2.2 各维度详细评分
    p = doc.add_paragraph()
    p.add_run('2.2 各维度详细评分').bold = True
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['维度', '均分(/5)', '有效样本', '低分率(≤2)', '备注']):
        hdr[i].text = h
    for dim in model['dimensions']:
        row = table.add_row().cells
        row[0].text = dim['label']
        row[1].text = f'{dim["mean"]:.2f}'
        row[2].text = str(dim['count'])
        row[3].text = f'{dim["low_rate"]}%'
        row[4].text = dim['note']
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f'总体对比（有效{ovs["total"]}条）：').bold = True
    if ovs['total']:
        p.add_run(f'AI优={ovs["win"]}({ovs["win_rate"]}%), 持平={ovs["draw"]}({ovs["draw_rate"]}%), AI劣={ovs["lose"]}({ovs["lose_rate"]}%)')

    # 2.3 各业务类型关键维度得分（chart1）
    p = doc.add_paragraph()
    p.add_run('2.3 各业务类型关键维度得分').bold = True
    if chart1_buf:
        doc.add_picture(chart1_buf, width=Inches(6.0))
    doc.add_paragraph()

    # 2.4 AI搜索 vs 文本搜索对比（chart2）
    p = doc.add_paragraph()
    p.add_run('2.4 AI搜索 vs 文本搜索对比').bold = True
    if chart2_buf:
        doc.add_picture(chart2_buf, width=Inches(5.5))
    doc.add_paragraph()

    # 2.5 按query特征切分
    p = doc.add_paragraph()
    p.add_run('2.5 按query特征切分').bold = True
    feat = doc.add_table(rows=1, cols=6)
    feat.style = 'Table Grid'
    fh = feat.rows[0].cells
    for i, h in enumerate(['切分维度', '分组', 'AI综合分', '相关性', '排序合理性', '质量门槛']):
        fh[i].text = h
    prev_dim = None
    for ft in model['features']:
        row = feat.add_row().cells
        row[0].text = ft['dim'] if ft['dim'] != prev_dim else ''
        prev_dim = ft['dim']
        row[1].text = f'{ft["group"]}(n={ft["n"]})'
        row[2].text = f'{ft["ai"]:.2f}' if ft['ai'] is not None else '-'
        row[3].text = f'{ft["relevance"]:.2f}' if ft['relevance'] is not None else '-'
        row[4].text = f'{ft["ranking"]:.2f}' if ft['ranking'] is not None else '-'
        row[5].text = f'{ft["quality"]:.2f}' if ft['quality'] is not None else '-'
    doc.add_paragraph()

    # 2.6 分业务类型表现
    p = doc.add_paragraph()
    p.add_run('2.6 分业务类型表现').bold = True
    ct = doc.add_table(rows=1, cols=6)
    ct.style = 'Table Grid'
    ch = ct.rows[0].cells
    for i, h in enumerate(['业务类型', 'query数量', '相关性', '排序合理性', '质量门槛', 'vs文本搜']):
        ch[i].text = h
    for d in categories:
        row = ct.add_row().cells
        row[0].text = d['cat']
        row[1].text = str(d['n'])
        row[2].text = f'{d["relevance"]:.2f}'
        row[3].text = f'{d["ranking"]:.2f}'
        row[4].text = f'{d["quality"]:.2f}'
        row[5].text = f'{d["win"]}胜{d["draw"]}平{d["lose"]}负' if d['vs_n'] else '-'

    # === 四、bad case梳理 ===
    doc.add_heading('四、bad case梳理', level=1)
    if notes.get('badcase', '').strip():
        render_badcase_html(doc, notes['badcase'])
    else:
        p = doc.add_paragraph()
        p.add_run('4.1 问题汇总').bold = True
        bc = doc.add_table(rows=1, cols=4)
        bc.style = 'Table Grid'
        bh = bc.rows[0].cells
        for i, h in enumerate(['体感', '问题', '原因', '解决办法']):
            bh[i].text = h
        for _ in range(3):
            bc.add_row()
        doc.add_paragraph('（按业务类型分节梳理典型bad case，请手动补充query/期望结果/实际问题/可能原因）')

    # === 附：评测方法 ===
    doc.add_heading('附：评测方法', level=1)
    doc.add_paragraph(f'• 样本：真实搜索日志按业务类型、频率分层抽样{total}条')
    doc.add_paragraph('• 评分：4维度×5分制（相关性、排序合理性、多样性、质量门槛），人工逐条评估')
    doc.add_paragraph('• 对比：同query同时执行文本搜索和AI搜索，对比打分')
    doc.add_paragraph('• 参考标准：Google Search Quality Rater Guidelines (Needs Met)')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=AI_search_report.docx"}
    )


# ============================================================
# Custom Columns API
# ============================================================

@app.get("/api/custom-columns")
async def get_custom_columns(test_set_id: int = 1):
    return load_custom_columns(test_set_id)


@app.post("/api/custom-columns")
async def add_custom_column(body: dict):
    label = body.get("label", "").strip()
    col_type = body.get("type", "text")  # "text" or "number"
    test_set_id = body.get("test_set_id", 1)
    if not label:
        raise HTTPException(400, "列名不能为空")

    # Generate field name from label (safe identifier)
    field = "custom_" + label.replace(" ", "_").replace("/", "_").replace("（", "_").replace("）", "_")
    # Ensure unique
    cols = load_custom_columns(test_set_id)
    existing_fields = {c["field"] for c in cols}
    existing_labels = {c["label"] for c in cols}
    if label in existing_labels:
        raise HTTPException(400, f"列 '{label}' 已存在")
    if field in existing_fields:
        field = field + "_" + str(len(cols))

    new_col = {"field": field, "label": label, "type": col_type}
    cols.append(new_col)
    save_custom_columns(cols, test_set_id)

    # ALTER TABLE (column is shared in DB, just visibility differs per test set)
    col_type_sql = "REAL" if col_type == "number" else "TEXT"
    async with aiosqlite.connect(DB_PATH) as db:
        try:
            await db.execute(f'ALTER TABLE queries ADD COLUMN {field} {col_type_sql}')
            await db.commit()
        except:
            pass  # column already exists

    return new_col


@app.delete("/api/custom-columns/{field}")
async def delete_custom_column(field: str, test_set_id: int = 1):
    cols = load_custom_columns(test_set_id)
    cols = [c for c in cols if c["field"] != field]
    save_custom_columns(cols, test_set_id)
    return {"ok": True}


# ============================================================
# Test Automation
# ============================================================

import subprocess
import threading
import signal

PROGRESS_FILE = os.path.join(os.path.dirname(APP_DIR), "search_progress.txt")
test_state = {"running": False, "progress": [], "total": 0, "current": 0, "status": "idle", "pid": None}


def read_breakpoint():
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, "r") as f:
            content = f.read().strip()
            return int(content) if content else 0
    return 0


def write_breakpoint(val):
    with open(PROGRESS_FILE, "w") as f:
        f.write(str(val))


@app.get("/api/test/status")
async def get_test_status():
    return {
        **test_state,
        "breakpoint": read_breakpoint()
    }


@app.post("/api/test/start")
async def start_test(body: dict = {}):
    if test_state["running"]:
        raise HTTPException(400, "测试正在运行中")

    # Allow overriding breakpoint
    bp = body.get("breakpoint")
    if bp is not None:
        write_breakpoint(int(bp))

    test_set_id = body.get("test_set_id", 1)

    # 开始标注即解锁该测试集
    db = await get_db()
    await db.execute("UPDATE test_sets SET status = 'labeling' WHERE id = ?", [test_set_id])
    await db.commit()
    await db.close()

    test_state["running"] = True
    test_state["progress"] = []
    test_state["status"] = "starting"
    test_state["current"] = 0
    test_state["total"] = 0

    def run_test():
        try:
            script_path = os.path.join(os.path.dirname(APP_DIR), "search_auto_test.py")
            env = os.environ.copy()
            env["TEST_SET_ID"] = str(test_set_id)
            proc = subprocess.Popen(
                ["python", "-u", script_path],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                cwd=os.path.dirname(APP_DIR), text=True, encoding="utf-8", errors="replace",
                env=env
            )
            test_state["pid"] = proc.pid
            for line in proc.stdout:
                line = line.strip()
                if line:
                    test_state["progress"].append(line)
                    # Keep last 200 lines
                    if len(test_state["progress"]) > 200:
                        test_state["progress"] = test_state["progress"][-200:]
                    if line.startswith("[") and "/" in line.split("]")[0]:
                        try:
                            parts = line.split("]")[0].strip("[")
                            cur, tot = parts.split("/")
                            test_state["current"] = int(cur)
                            test_state["total"] = int(tot)
                        except:
                            pass
                    test_state["status"] = "running"
            proc.wait()
            test_state["status"] = "done" if proc.returncode == 0 else "error"
        except Exception as e:
            test_state["progress"].append(f"启动失败: {e}")
            test_state["status"] = "error"
        finally:
            test_state["running"] = False
            test_state["pid"] = None

    threading.Thread(target=run_test, daemon=True).start()
    return {"ok": True, "message": "测试已启动"}


@app.post("/api/test/stop")
async def stop_test():
    if test_state["pid"]:
        try:
            os.kill(test_state["pid"], signal.SIGTERM)
        except:
            pass
    test_state["running"] = False
    test_state["status"] = "stopped"
    test_state["pid"] = None
    return {"ok": True}


@app.post("/api/test/breakpoint")
async def set_breakpoint(body: dict):
    val = body.get("value", 0)
    write_breakpoint(int(val))
    return {"ok": True, "breakpoint": int(val)}


@app.post("/api/test/finish")
async def finish_test(body: dict = {}):
    """结束标注：停止脚本 + 清空断点 + 锁定测试集。"""
    if test_state["pid"]:
        try:
            os.kill(test_state["pid"], signal.SIGTERM)
        except Exception:
            pass
    test_state["running"] = False
    test_state["status"] = "finished"
    test_state["pid"] = None
    write_breakpoint(0)  # 清空断点

    test_set_id = body.get("test_set_id")
    if test_set_id:
        db = await get_db()
        await db.execute("UPDATE test_sets SET status = 'locked' WHERE id = ?", [test_set_id])
        await db.commit()
        await db.close()
    return {"ok": True, "locked": test_set_id}


# ============================================================
# Frontend
# ============================================================

@app.get("/", response_class=HTMLResponse)
async def index():
    html_path = os.path.join(STATIC_DIR, "index.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/analysis", response_class=HTMLResponse)
async def analysis_page():
    html_path = os.path.join(STATIC_DIR, "analysis.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/detail", response_class=HTMLResponse)
async def detail_page():
    html_path = os.path.join(STATIC_DIR, "detail.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/annotate", response_class=HTMLResponse)
async def annotate_page():
    html_path = os.path.join(STATIC_DIR, "annotate.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/knowledge", response_class=HTMLResponse)
async def knowledge_page():
    html_path = os.path.join(STATIC_DIR, "knowledge.html")
    with open(html_path, "r", encoding="utf-8") as f:
        return f.read()


@app.get("/api/knowledge")
async def get_knowledge():
    return load_knowledge()


@app.put("/api/knowledge")
async def put_knowledge(body: dict):
    return save_knowledge(body)


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)
